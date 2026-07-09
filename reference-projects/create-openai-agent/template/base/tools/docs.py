"""
Built-in document generation tools.
Supports: Word (.docx), Excel (.xlsx), Markdown → PDF.

Install: uv add 'create-agent[docs]'  or  uv add python-docx openpyxl
"""

from __future__ import annotations
from pathlib import Path
from agents import function_tool


@function_tool
def create_word_doc(file_path: str, content: str, title: str = "") -> str:
    """
    Create a Word document (.docx) from markdown-style content.
    Lines starting with '# ' become Heading 1, '## ' → Heading 2, etc.
    All other lines become paragraphs.
    Returns the absolute path of the created file.
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: uv add 'create-agent[docs]'"

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    path = Path(file_path).expanduser()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return f"Created: {path.resolve()}"


@function_tool
def create_excel(file_path: str, data: str, sheet_name: str = "Sheet1") -> str:
    """
    Create an Excel file (.xlsx) from CSV-formatted data.
    data: CSV text (first row = header).
    Returns the absolute path of the created file.

    Example data:
      Name,Age,City
      Alice,30,Hong Kong
      Bob,25,London
    """
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl not installed. Run: uv add 'create-agent[docs]'"

    import csv, io
    rows = list(csv.reader(io.StringIO(data.strip())))
    if not rows:
        return "Error: no data provided"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name

    for row in rows:
        ws.append(row)

    # Auto-width columns
    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

    path = Path(file_path).expanduser()
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    return f"Created: {path.resolve()}"


@function_tool
def markdown_to_pdf(file_path: str, markdown_content: str) -> str:
    """
    Convert markdown text to a PDF file.
    Returns the absolute path of the created file.
    Install: uv add 'create-agent[docs]'  (requires weasyprint or md-to-pdf)
    """
    try:
        import markdown as md_lib
        from weasyprint import HTML
    except ImportError:
        return (
            "Error: weasyprint not installed. Run: uv add 'create-agent[docs]'\n"
            "On macOS: brew install weasyprint"
        )

    html_body = md_lib.markdown(markdown_content, extensions=["tables", "fenced_code"])
    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: -apple-system, sans-serif; max-width: 800px; margin: 40px auto; line-height: 1.6; color: #222; }}
  h1,h2,h3 {{ color: #111; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
  code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
  pre {{ background: #f4f4f4; padding: 16px; border-radius: 6px; overflow-x: auto; }}
</style></head>
<body>{html_body}</body></html>"""

    path = Path(file_path).expanduser()
    path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html).write_pdf(str(path))
    return f"Created: {path.resolve()}"


# Convenience bundle
DOCS_TOOLS = [create_word_doc, create_excel, markdown_to_pdf]
